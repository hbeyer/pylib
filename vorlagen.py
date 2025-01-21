#!/usr/bin/python3
# -*- coding: utf-8 -*-

"""
# Eine Liste (geordnet, indexiert und veränderlich)
mylist = ['Lerche', 'Schneider', 'Zimmermann', 'Kästner', 'Raabe', 'Schmidt-Glintzer', 'bURSCHEL']
mylist[len(mylist) - 1] = mylist[len(mylist) - 1].swapcase()
mylist.append('Ritter Rost')
mylist.insert(0, 'Zimmermann')
print(mylist)
"""

"""
# Iterieren über eine Liste mit Zugriff auf die Indices
mylist = ['Lerche', 'Schneider', 'Zimmermann', 'Kästner', 'Raabe', 'Schmidt-Glintzer', 'Burschel']
for ind, val in enumerate(mylist):
    print(f"{ind}: {val}")
"""

"""
# Ein Tupel (ist unveränderlich)
mytuple = ('Montag', 'Dienstag', 'Mittwoch', 'Donnerstag', 'Freitag', 'Samstag', 'Sonntag')
#print(mytuple[3:6])
"""

"""
# Unpacking eines Tupel
mytuple = ('Montag', 'Dienstag', 'Mittwoch', 'Donnerstag', 'Freitag', 'Samstag', 'Sonntag')
Mohn, Dienst, Mitte, Donner, Frei, Sams, Sonne = mytuple
print(f"Das Sams kommt an diesem Tag: {Sams}")
"""

"""
# Ein Set (unindexiert und ungeordnet, Elemente sind unveränderlich, können aber vermehrt oder reduziert werden)
myset = {'Adenauer', 'Erhard', 'Kiesinger', 'Brandt', 'Schmidt', 'Kohl', 'Schröder', 'Merkel', 'Schulz'}
myset.remove('Schulz')
myset.add('Kramp-Karrenbauer')
for i in myset:
    print(i)
"""

"""
# Ein Dictionary
mydict = {'Mann':'vyras', 'Frau':'moteris','Fisch':'žuvis', 'Biber':'bebras', 'Stadt':'miestas', 'König':'karalius'}
for x, y in mydict.items():
  print(f"{x} heißt auf Litauisch {y}")
"""
    
"""
# Eine Datumsoperation
import time
import datetime

time = time.localtime(time.time())
print(time)
"""

"""
# Eine Funktion
def makeName(forename, surname, title=""):
	result = forename + " " + surname
	if title:
		result = title + " " + result
	return result

print(makeName("Hartmut", "Beyer", "Magister artium"))
"""

"""
# Eine anonyme Funktion
x = lambda e, f: str(e) + " - " + str(f)
print(x(1758, "Hochkirch"))
"""

"""
# Gruppieren einer Datensammlung nach einem darin enthaltenen Kriterium:

import itertools
from operator import itemgetter

books = [
	{"author":"Luther, Martin", "title":"Ein Sendbrief von dem harten Büchlein wider die Bauren", "year":"1525"},
	{"author":"Luther, Martin", "title":"An die Pfarrherrn, wider den Wucher zu predigen", "year":"1540"},
	{"author":"Luther, Martin", "title":"Vom Ehebruch vnd Weglauffen", "year":"1540"} ,   
	{"author":"Melanchthon, Philipp", "title":"De Officio Principum", "year":"1540"},
	{"author":"Melanchthon, Philipp", "title":"Commentaria in Epistolam Pauli ad Col.", "year":"1547"},
	{"author":"Luther, Martin", "title":"Ettlich Artickelstuck so Mart.Luther erhalten wil wider die gantze Satans schuole", "year":"1530"},    
	{"author":"Bugenhagen, Johannes", "title":"Wie es uns zu Wittenberg in der Stadt gegangen ist in diesem vergangen Krieg", "year":"1540"}    
]
sorted_books = sorted(books, key=itemgetter("year"))
groups = itertools.groupby(sorted_books, key=lambda x:x["year"])

for group in groups:
	#print(group[0] + ": " + "/".join([row["author"] for row in group[1]]))
        print(f'{group[0]}: {"/".join([row["author"] for row in group[1]])}')
"""

"""
# Eine Klasse
class Person:
	def __init__(self, forename, surname):
		self.forename = forename
		self.surename = surname

person = Person('David', 'Ben-Gurion')
print(person.forename)
"""

"""
# Eine Klasse
class Language:
	def __init__(self,code):
		self.codes = {
			"eng":"Englisch", 
			"ger":"Deutsch", 
			"fre":"Französisch",
			"rus":"Russisch"
		}
		if code not in self.codes:
			self.name = code
			return
		self.name = self.codes[code]		
	
lang = Language("rus")
print(lang.name)
"""

"""
# Eine Datei aus dem Netz auslesen
import urllib.request as ur
url = "http://diglib.hab.de/edoc/ed000228/1623_06.xml"
fileobject = ur.urlopen(url)
string = fileobject.read()
print(string)
"""

"""
# Einlesen einer Datei, Durchsuchen mit einem regulären Ausdruck und Wiedergabe des Ergebnisses
import re
with open("test.txt", "r") as file:
    istc = re.findall(r"ISTC (.+)", file.read())
    print('|'.join(istc))
"""

"""
# Eine XML-Datei parsen
import xml.etree.ElementTree as et
tree = et.parse('test.xml')
root = tree.getroot()
nbs = root.findall('.//{http://www.tei-c.org/ns/1.0}rs')
name = ""
for ent in nbs:
    if ent.get('type') == 'person':
        name = str(ent.text).strip()
        ref = str(ent.get('ref')).strip()
        print(name + ' - ' + ref)
"""

"""
# Laden und Auslesen einer XML-Datei im Netz
import urllib.request as ur
import xml.etree.ElementTree as et

url = "http://diglib.hab.de/edoc/ed000228/1623_08.xml"
fileobject = ur.urlopen(url)
tree = et.parse(fileobject)
root = tree.getroot()
nbs = root.findall('.//{http://www.tei-c.org/ns/1.0}rs')
name = ""
for ent in nbs:
    if ent.get('type') == 'person':
        name = str(ent.text).strip()
        ref = str(ent.get('ref')).strip()
        print(name + ' - ' + ref)
"""

"""
# Verändern einer XML-Datei

import xml.etree.ElementTree as et

project = 'caselius'
pathInputFile = 'projectFiles/' + project + '/' + project + '-pod.xml'
pathOutputFile = '' + project + '-bearb.xml'
separator = ';'

tree = et.parse(pathInputFile)
root = tree.getroot()

items = root.findall('.//item')
for element in items:
	publishers = element.findall('publisher')
	if publishers != None:
		for publisher in publishers:
			pubList = et.SubElement(element, 'publishers')
			splitPublishers = publisher.text.split(separator)
			for pub in splitPublishers:
				newPub = et.SubElement(pubList, 'publisher')
				newPub.text = pub.strip()
			element.remove(publisher)

tree.write(pathOutputFile, encoding="UTF-8", xml_declaration=True)
"""

"""
# Eine Datenbankabfrage mit List Comprehension und Einsatz der Funktion join
import mysql.connector

mydb = mysql.connector.connect(
  host = "localhost",
  user = "root",
  passwd = "{Passwort}",
  database = "{Datenbank}"
)

mycursor = mydb.cursor()
mycursor.execute("SELECT id, gnd FROM helmstedt.temp_prof_kat")
myresult = mycursor.fetchall()
gnds = [x[1] for x in myresult if x[1] != None]
print('|'.join(gnds))
"""

"""
# Nutzen einer internen SQLite-Datenbank

# Anlegen der Datenbank
import sqlite3
conn = sqlite3.connect("database.db")
cursor = conn.cursor()
cursor.execute('''CREATE TABLE records (id int, bbg text, vdn text, vd16m text, title text, jahr text)''')
conn.commit()
conn.close()

# Befüllen der Datenbank
conn = sqlite3.connect("database.db")
cursor = conn.cursor()

val = [0, "Aav", "VD16 ZV 22046", "VD0053034", "Die @Wunderlichst vnerhörtest Legend || Vnd beschreibung || Des abgeführten Quartirten Ge=||vierten vnd Viereckechten Vierhörnigen Hütleins", "1603"]
cursor.execute("INSERT INTO records VALUES (?, ?, ?, ?, ?, ?)", val)
conn.commit()
conn.close()

# Abfragen der Datenbank mit einem Context Manager
with sqlite3.connect("database.db") as conn:
    cursor = conn.cursor()
    sql = f"SELECT * FROM main"
    cursor.execute(sql)
    result = cursor.fetchall()
    print(result)
"""

"""
# Übertragen von Daten aus einer CSV-Tabelle in eine interne Datenbank

conn = sqlite3.connect("vd218.db")
cursor = conn.cursor()
file = open('vd218.csv', newline='')
reader = csv.reader(file, delimiter=',', quotechar='|')
count = 0
for line in reader:
	val = [count, line[0], line[1], None, None]
	cursor.execute("INSERT INTO records VALUES (?, ?, ?, ?, ?)", val)
	count += 1
conn.commit()
conn.close()
"""

"""
# Ablage von Programmdaten in einer Datei

import pickle
ser = ["Himbeer", "Zitrone", "Ananas"]
pickle.dump(ser, open("obst.p", "wb"))
obst = pickle.load(open("obst.p", "rb"))
print(obst)
"""

"""
# Einen Logger betreiben

import logging

# Das Folgende sorgt dafür, dass ab diesem Level Meldungen angezeigt werden
# Sollte da stehen, wo der Code ausgeführt wird (u. U. main() nutzen)

logging.basicConfig()
logging.getLogger().setLevel(logging.DEBUG)

logging.debug("Die Checksumme beträgt 45678")
logging.info("Es wurden alle Dateien geladen")
logging.warning("Laden der nicht essenziellen Datei xy fehlgeschlagen")
logging.error("Es ist ein Problem aufgetreten")
logging.critical("Wegen eines Problems muss das Programm beendet werden")
"""

"""
# Erstellem eines HTML-Dokuments

from bs4 import BeautifulSoup

def makeOPACLink(shelfmark):
    shelfmark = shelfmark.replace("(", "").replace(")", "")
    shelfmark = urllib.parse.quote_plus(shelfmark)
    return(f"http://opac.lbs-braunschweig.gbv.de/DB=2/CMD?ACT=SRCHA&TRM=sgb+{shelfmark}")

htmbase = f"<!DOCTYPE html><html><head><meta charset=\"UTF-8\"></head><body><h1>Neuerscheinungen der Sammlung Deutscher Drucke 1601&ndash;1700</h1></body></html>"
soup = BeautifulSoup(htmbase, 'html.parser')

orig_tag = soup.h1
h2 = soup.new_tag("h2")
h2.string = "02 Philosophie"
orig_tag.insert_after(h2)
orig_tag = h2
par = soup.new_tag("p")
par.string = "Crell, Fortunatus: Fortunati Crellii Isagoge Logica. In Duas Partes Tributa ; In communem Et Propriam. Editio Septima. Neustadt an der Weinstraße: Schramm, Nikolaus, Wilhelm Harnisch Erben 1602 "
ahref = soup.new_tag("a")
ahref.string = cp.sm
ahref["href"] = makeOPACLink(cp.sm)
ahref["target"] = "_blank"
ahref["title"] = "Titel im OPAC anzeigen"
par.append(ahref)
orig_tag.insert_after(par)

with open("SDDList.htm", "w", encoding="utf-8") as file:
    file.write(str(soup.prettify()))
"""

"""
# Erstellen oder Verändern eines Word-Dokuments
from docx import Document
doc = Document("Vorlage.docx")
# Leeres Dokument:
# doc = Document()
doc.add_heading(f"Titel des Dokuments", level=1)
doc.add_heading("Untertitel", level=2)
doc.add_paragraph("Das ist ein Absatz.")
doc.save(f"Mein_Word-Dokument.docx")
"""